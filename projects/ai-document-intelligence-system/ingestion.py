"""
ingestion.py
============
Document Intelligence System — Multi-Format Ingestion
Project: P9 · prompt-engineering-lab

Supported formats:
  .txt   — plain text (native)
  .csv   — tabular data (native)
  .pdf   — via pypdf (fallback: raw bytes warning)
  .docx  — via python-docx (fallback: zip extraction)
  .md    — markdown treated as plain text

Returns a DocumentRaw object with text content + metadata.
"""

import re
import logging
from pathlib import Path
from dataclasses import dataclass, field
from typing import Optional

logger = logging.getLogger(__name__)


@dataclass
class DocumentRaw:
    """Raw document content plus file metadata."""
    path: str
    filename: str
    extension: str
    text: str
    page_count: int = 1
    word_count: int = 0
    char_count: int = 0
    metadata: dict = field(default_factory=dict)
    extraction_method: str = "native"
    error: Optional[str] = None

    def __post_init__(self):
        if self.text:
            self.word_count = len(re.findall(r'\b\w+\b', self.text))
            self.char_count = len(self.text)


def ingest(file_path: str) -> DocumentRaw:
    """
    Ingest a document from any supported format.
    Returns DocumentRaw with full text content.

    Args:
        file_path: Path to the document file

    Returns:
        DocumentRaw with extracted text and metadata
    """
    path = Path(file_path)
    ext  = path.suffix.lower()

    if not path.exists():
        return DocumentRaw(
            path=str(path), filename=path.name, extension=ext,
            text="", error=f"File not found: {file_path}"
        )

    dispatcher = {
        ".txt":  _ingest_txt,
        ".md":   _ingest_txt,
        ".csv":  _ingest_csv,
        ".pdf":  _ingest_pdf,
        ".docx": _ingest_docx,
    }

    handler = dispatcher.get(ext, _ingest_txt)
    return handler(path)


# ── Format handlers ──────────────────────────────────────────

def _ingest_txt(path: Path) -> DocumentRaw:
    try:
        text = path.read_text(encoding="utf-8", errors="replace")
        return DocumentRaw(
            path=str(path), filename=path.name,
            extension=path.suffix.lower(), text=text,
            extraction_method="native_text",
        )
    except Exception as e:
        return DocumentRaw(
            path=str(path), filename=path.name,
            extension=path.suffix.lower(), text="", error=str(e),
        )


def _ingest_csv(path: Path) -> DocumentRaw:
    """Convert CSV to readable text representation."""
    try:
        import csv
        rows = []
        with open(path, encoding="utf-8", errors="replace") as f:
            reader = csv.DictReader(f)
            headers = reader.fieldnames or []
            rows.append("Columns: " + ", ".join(headers))
            rows.append("-" * 40)
            for i, row in enumerate(reader):
                row_text = " | ".join(f"{k}: {v}" for k, v in row.items())
                rows.append(f"Row {i+1}: {row_text}")

        text = "\n".join(rows)
        return DocumentRaw(
            path=str(path), filename=path.name, extension=".csv",
            text=text, extraction_method="csv_reader",
            metadata={"columns": headers, "row_count": len(rows) - 2},
        )
    except Exception as e:
        return _ingest_txt(path)


def _ingest_pdf(path: Path) -> DocumentRaw:
    """Extract text from PDF using pypdf, with raw fallback."""
    # Try pypdf
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        pages  = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                pages.append(f"[Page {i+1}]\n{text}")

        return DocumentRaw(
            path=str(path), filename=path.name, extension=".pdf",
            text="\n\n".join(pages),
            page_count=len(reader.pages),
            extraction_method="pypdf",
            metadata={"page_count": len(reader.pages)},
        )
    except ImportError:
        logger.warning("pypdf not installed. Install with: pip install pypdf")
    except Exception as e:
        logger.warning(f"pypdf extraction failed: {e}")

    # Fallback: try pdfminer
    try:
        from pdfminer.high_level import extract_text as pm_extract
        text = pm_extract(str(path))
        return DocumentRaw(
            path=str(path), filename=path.name, extension=".pdf",
            text=text or "", extraction_method="pdfminer",
        )
    except ImportError:
        pass
    except Exception as e:
        logger.warning(f"pdfminer extraction failed: {e}")

    return DocumentRaw(
        path=str(path), filename=path.name, extension=".pdf",
        text="", error="No PDF extraction library available. Install pypdf: pip install pypdf",
    )


def _ingest_docx(path: Path) -> DocumentRaw:
    """Extract text from DOCX using python-docx, with zip fallback."""
    try:
        from docx import Document
        doc    = Document(str(path))
        blocks = []
        for para in doc.paragraphs:
            if para.text.strip():
                style = para.style.name if para.style else ""
                if "Heading" in style:
                    blocks.append(f"\n## {para.text}")
                else:
                    blocks.append(para.text)

        # Extract tables
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                blocks.append("\n[TABLE]\n" + "\n".join(rows))

        return DocumentRaw(
            path=str(path), filename=path.name, extension=".docx",
            text="\n".join(blocks),
            extraction_method="python_docx",
            metadata={"paragraph_count": len(doc.paragraphs)},
        )
    except ImportError:
        logger.warning("python-docx not installed. Install with: pip install python-docx")
    except Exception as e:
        logger.warning(f"python-docx failed: {e}")

    # Fallback: extract XML from zip
    try:
        import zipfile
        with zipfile.ZipFile(str(path)) as z:
            with z.open("word/document.xml") as f:
                xml = f.read().decode("utf-8", errors="replace")
        text = re.sub(r'<[^>]+>', ' ', xml)
        text = re.sub(r'\s+', ' ', text).strip()
        return DocumentRaw(
            path=str(path), filename=path.name, extension=".docx",
            text=text, extraction_method="zip_xml_fallback",
        )
    except Exception as e:
        return DocumentRaw(
            path=str(path), filename=path.name, extension=".docx",
            text="", error=str(e),
        )


def ingest_directory(dir_path: str, extensions: list = None) -> list:
    """
    Ingest all supported documents in a directory.

    Args:
        dir_path:   Path to directory
        extensions: List of extensions to include (default: all supported)

    Returns:
        List of DocumentRaw objects
    """
    supported = extensions or [".txt", ".pdf", ".docx", ".csv", ".md"]
    results   = []
    for path in sorted(Path(dir_path).iterdir()):
        if path.suffix.lower() in supported:
            logger.info(f"  Ingesting: {path.name}")
            doc = ingest(str(path))
            results.append(doc)
    return results


if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)
    docs = ingest_directory("data/documents")
    for doc in docs:
        status = "✓" if not doc.error else "✗"
        print(f"{status} {doc.filename:45s}  {doc.word_count:5d} words  [{doc.extraction_method}]")
